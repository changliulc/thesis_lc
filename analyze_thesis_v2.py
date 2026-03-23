import os
from pypdf import PdfReader
from docx import Document
import re

def extract_pdf_text(pdf_path, max_pages=20):
    """提取PDF前N页文本"""
    try:
        reader = PdfReader(pdf_path)
        text = ""
        total_pages = len(reader.pages)
        pages_to_read = min(max_pages, total_pages)
        
        for i in range(pages_to_read):
            page_text = reader.pages[i].extract_text() or ""
            text += page_text + "\n"
        
        return text, total_pages
    except Exception as e:
        return f"Error reading PDF: {e}", 0

def extract_docx_text(docx_path):
    """提取docx文本"""
    try:
        doc = Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text, len(doc.paragraphs)
    except Exception as e:
        return f"Error reading DOCX: {e}", 0

def analyze_paper_quality(name, text, total_pages_or_paras):
    """分析论文质量指标"""
    
    # 基础统计
    total_chars = len(text)
    total_words = len(text.split())
    
    # 查找关键章节
    chapters = {
        "绪论/引言": bool(re.search(r'(绪论|引言|研究背景)', text)),
        "相关工作": bool(re.search(r'(相关工作|文献综述|国内外研究现状)', text)),
        "方法/算法": bool(re.search(r'(方法|算法|模型|系统设计)', text)),
        "实验": bool(re.search(r'(实验|仿真|测试|验证)', text)),
        "结论": bool(re.search(r'(结论|总结|展望)', text))
    }
    
    # 查找图表
    figures = len(re.findall(r'(图\s*\d+|Figure\s*\d+|Fig\.\s*\d+)', text, re.IGNORECASE))
    tables = len(re.findall(r'(表\s*\d+|Table\s*\d+)', text, re.IGNORECASE))
    
    # 查找公式
    equations = len(re.findall(r'(\$[^$]+\$|\\begin\{equation\}|\\\(|\\\[)', text))
    
    # 查找参考文献引用
    citations = len(re.findall(r'\\cite\{|\[\d+\]|\(\d+\)', text))
    
    # 查找英文摘要
    has_english_abstract = bool(re.search(r'(ABSTRACT|Abstract)', text))
    
    # 查找关键词
    has_keywords = bool(re.search(r'(关键词|Key\s*words)', text, re.IGNORECASE))
    
    return {
        "name": name,
        "total_chars": total_chars,
        "total_words": total_words,
        "pages_or_paras": total_pages_or_paras,
        "chapters": chapters,
        "figures": figures,
        "tables": tables,
        "equations": equations,
        "citations": citations,
        "has_english_abstract": has_english_abstract,
        "has_keywords": has_keywords
    }

def print_analysis(result):
    """打印分析结果"""
    print(f"\n{'='*80}")
    print(f"【{result['name']}】论文质量分析")
    print(f"{'='*80}")
    
    print(f"\n📊 基础统计:")
    print(f"  - 总字符数: {result['total_chars']:,}")
    print(f"  - 总词数: {result['total_words']:,}")
    print(f"  - 页数/段落数: {result['pages_or_paras']}")
    
    print(f"\n📑 章节完整性:")
    for chapter, exists in result['chapters'].items():
        status = "✓" if exists else "✗"
        print(f"  {status} {chapter}")
    
    print(f"\n📈 图表公式:")
    print(f"  - 图表数量: {result['figures']}")
    print(f"  - 表格数量: {result['tables']}")
    print(f"  - 公式数量: {result['equations']}")
    print(f"  - 引用数量: {result['citations']}")
    
    print(f"\n🌍 其他:")
    print(f"  - 英文摘要: {'✓' if result['has_english_abstract'] else '✗'}")
    print(f"  - 关键词: {'✓' if result['has_keywords'] else '✗'}")
    
    # 计算综合得分
    score = 0
    score += min(result['total_chars'] / 10000, 10)  # 字数得分
    score += sum(result['chapters'].values()) * 5  # 章节完整性
    score += min(result['figures'] * 2, 10)  # 图表
    score += min(result['tables'] * 2, 10)  # 表格
    score += min(result['equations'], 10)  # 公式
    score += min(result['citations'], 10)  # 引用
    score += 5 if result['has_english_abstract'] else 0
    score += 5 if result['has_keywords'] else 0
    
    print(f"\n⭐ 综合得分: {score:.1f}/100")
    
    return score

# 文件路径
base_path = r"D:\my_app\wechat\xwechat_files\wxid_4ns7hqjaxq2l22_1cbc\msg\file\2026-03"

files = {
    "侯照莹": os.path.join(base_path, "侯照莹_面向双车道高速公路的多地磁传感器车辆目标跟踪方法研究_盲审版.pdf"),
    "黄研": os.path.join(base_path, "黄研-低功耗抗干扰地磁检测算法研究0320.docx"),
    "刘畅": os.path.join(base_path, "刘畅_面向智慧交通的车型分类及停车占用检测方法研究_盲审版.pdf"),
    "邱嘉乐": os.path.join(base_path, "邱嘉乐_面向智慧公路的LoRa网关设计与组网优化研究.pdf")
}

print("=" * 80)
print("四篇硕士学位论文质量分析")
print("=" * 80)

scores = {}

for name, path in files.items():
    print(f"\n正在分析: {name}...")
    
    if not os.path.exists(path):
        print(f"❌ 文件不存在: {path}")
        continue
    
    if path.endswith('.pdf'):
        text, total = extract_pdf_text(path, max_pages=30)
    elif path.endswith('.docx'):
        text, total = extract_docx_text(path)
    else:
        print(f"❌ 不支持的文件格式")
        continue
    
    result = analyze_paper_quality(name, text, total)
    score = print_analysis(result)
    scores[name] = score

# 排名
print(f"\n{'='*80}")
print("🏆 最终排名")
print(f"{'='*80}")

ranked = sorted(scores.items(), key=lambda x: x[1], reverse=True)
for i, (name, score) in enumerate(ranked, 1):
    medal = ["🥇", "🥈", "🥉", "4️⃣"][i-1]
    print(f"{medal} 第{i}名: {name} - {score:.1f}分")

print(f"\n{'='*80}")
