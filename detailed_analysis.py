import os
from pypdf import PdfReader
from docx import Document
import re

def extract_pdf_full_text(pdf_path):
    """提取完整PDF文本"""
    try:
        reader = PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n"
        return text, len(reader.pages)
    except Exception as e:
        return f"Error: {e}", 0

def analyze_content_quality(name, text, total_pages):
    """深度内容分析"""
    
    print(f"\n{'='*80}")
    print(f"【{name}】深度内容分析")
    print(f"{'='*80}")
    
    # 提取摘要部分
    abstract_match = re.search(r'摘\s*要(.*?)关\s*键\s*词', text, re.DOTALL | re.IGNORECASE)
    if abstract_match:
        abstract = abstract_match.group(1).strip()[:500]
        print(f"\n📝 摘要预览:")
        print(f"{abstract}...")
    
    # 查找研究内容
    research_content = re.findall(r'(研究内容|主要工作|本文主要|本文研究)(.*?)(?=(章节安排|论文结构|本文结构|$))', 
                                   text, re.DOTALL | re.IGNORECASE)
    if research_content:
        print(f"\n🔬 研究内容:")
        for _, content in research_content[:2]:
            clean_content = content.strip()[:300]
            print(f"  {clean_content}...")
    
    # 查找创新点
    innovation = re.findall(r'(创新点|主要贡献|创新之处)(.*?)(?=(研究内容|章节安排|$))', 
                           text, re.DOTALL | re.IGNORECASE)
    if innovation:
        print(f"\n💡 创新点:")
        for _, content in innovation[:1]:
            clean_content = content.strip()[:300]
            print(f"  {clean_content}...")
    
    # 统计各章篇幅
    chapters = re.findall(r'(第[一二三四五六]章|Chapter\s*\d+)(.*?)(?=(第[一二三四五六]章|Chapter\s*\d+|结论|总结|$))', 
                         text, re.DOTALL | re.IGNORECASE)
    print(f"\n📚 章节数量: {len(chapters)}")
    
    # 查找实验数据
    experiments = re.findall(r'(实验|仿真|测试|验证)(.*?)(?=(本章小结|结论|$))', 
                            text, re.DOTALL | re.IGNORECASE)
    print(f"🔬 实验/测试部分: {len(experiments)} 处")
    
    # 查找性能指标
    metrics = re.findall(r'(准确率|精度|召回率|F1|准确率|accuracy|precision|recall)', 
                        text, re.IGNORECASE)
    print(f"📊 性能指标提及: {len(metrics)} 次")
    
    # 查找算法/方法
    algorithms = re.findall(r'(算法|方法|模型|网络|系统)(.*?)(?=(设计|实现|提出))', 
                           text, re.DOTALL | re.IGNORECASE)
    print(f"⚙️ 算法/方法: {len(algorithms)} 处")
    
    # 技术关键词统计
    tech_keywords = {
        "深度学习": len(re.findall(r'深度学习|神经网络|CNN|RNN|LSTM|Transformer', text)),
        "机器学习": len(re.findall(r'机器学习|SVM|随机森林|决策树|聚类', text)),
        "信号处理": len(re.findall(r'滤波|FFT|频域|时域|信号处理', text)),
        "传感器": len(re.findall(r'传感器|地磁|雷达|视频|激光', text)),
        "通信": len(re.findall(r'LoRa|通信|传输|网络|协议', text)),
        "优化": len(re.findall(r'优化|遗传算法|粒子群|梯度下降', text))
    }
    
    print(f"\n🎯 技术关键词分布:")
    for keyword, count in tech_keywords.items():
        if count > 0:
            print(f"  - {keyword}: {count}次")
    
    return True

# 文件路径
base_path = r"D:\my_app\wechat\xwechat_files\wxid_4ns7hqjaxq2l22_1cbc\msg\file\2026-03"

files = {
    "侯照莹": os.path.join(base_path, "侯照莹_面向双车道高速公路的多地磁传感器车辆目标跟踪方法研究_盲审版.pdf"),
    "黄研": os.path.join(base_path, "黄研-低功耗抗干扰地磁检测算法研究0320.docx"),
    "刘畅": os.path.join(base_path, "刘畅_面向智慧交通的车型分类及停车占用检测方法研究_盲审版.pdf"),
    "邱嘉乐": os.path.join(base_path, "邱嘉乐_面向智慧公路的LoRa网关设计与组网优化研究.pdf")
}

print("=" * 80)
print("四篇硕士学位论文深度分析")
print("=" * 80)

for name, path in files.items():
    if not os.path.exists(path):
        print(f"\n❌ {name}: 文件不存在")
        continue
    
    try:
        if path.endswith('.pdf'):
            text, total = extract_pdf_full_text(path)
        elif path.endswith('.docx'):
            doc = Document(path)
            text = "\n".join([para.text for para in doc.paragraphs])
            total = len(doc.paragraphs)
        else:
            continue
        
        analyze_content_quality(name, text, total)
        
    except Exception as e:
        print(f"\n❌ {name}: 分析出错 - {e}")

print(f"\n{'='*80}")
print("分析完成")
print("=" * 80)
